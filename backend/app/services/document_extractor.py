"""
Document text extractor.
Supports PDF (pdfplumber → pypdf fallback), DOCX, TXT, MD.
Goal: extract ALL text with zero loss.
"""
import io
from pathlib import Path


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _pdf(content)
    elif ext in (".docx",):
        return _docx(content)
    elif ext in (".txt", ".md"):
        return content.decode("utf-8", errors="replace")
    else:
        # Try UTF-8 as a last resort
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")


def _pdf(content: bytes) -> str:
    """Extract text from PDF — tries pdfplumber first, falls back to pypdf."""
    text_parts = []

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(f"[Page {i + 1}]\n{t.strip()}")
        if text_parts:
            return "\n\n".join(text_parts)
    except Exception:
        pass  # Fall through to pypdf

    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        for i, page in enumerate(reader.pages):
            t = page.extract_text()
            if t and t.strip():
                text_parts.append(f"[Page {i + 1}]\n{t.strip()}")
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")


def _docx(content: bytes) -> str:
    """Extract text from DOCX preserving headings and paragraphs."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Mark headings
                if para.style and "Heading" in (para.style.name or ""):
                    lines.append(f"\n### {para.text.strip()}")
                else:
                    lines.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}")
